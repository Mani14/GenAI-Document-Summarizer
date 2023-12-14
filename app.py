from flask import Flask, render_template, request, jsonify, send_file
import os
from azure.storage.blob import BlobServiceClient
import json
from azure.data.tables import TableServiceClient
import docx
from docx.shared import Inches

# Get the current working directory
current_path = os.getcwd()
print(f"Current Path: {current_path}")

app = Flask(__name__)

# Initializing table service and table client
connection_string = "DefaultEndpointsProtocol=https;AccountName=gvdstorage;AccountKey=tam4CCOAGP2XxBWfylUNmx3isscihpX/tt7dKW4ZTQiYZUT8gIYmjDZ7FxFHar/z6dJpiGD1NE2H+AStmEmQiw==;EndpointSuffix=core.windows.net"
table_service_client = TableServiceClient.from_connection_string(conn_str=connection_string)

table_name = "GVDSummary"
table_client = table_service_client.get_table_client(table_name)

# Azure Storage account details
ACCOUNT_NAME = 'gvdstorage'
ACCOUNT_KEY = 'tam4CCOAGP2XxBWfylUNmx3isscihpX/tt7dKW4ZTQiYZUT8gIYmjDZ7FxFHar/z6dJpiGD1NE2H+AStmEmQiw'
CONTAINER_NAME = 'gvdcontainer2'

# Function to initialize the BlobServiceClient
def initialize_blob_service_client():
    connection_string = "BlobEndpoint=https://gvdstorage.blob.core.windows.net/;QueueEndpoint=https://gvdstorage.queue.core.windows.net/;FileEndpoint=https://gvdstorage.file.core.windows.net/;TableEndpoint=https://gvdstorage.table.core.windows.net/;SharedAccessSignature=sv=2022-11-02&ss=bfqt&srt=sco&sp=rwdlacupiytfx&se=2024-04-01T15:49:05Z&st=2023-12-11T07:49:05Z&spr=https,http&sig=9hLYSeyh%2B76Dx3rOUbNMhNT2jvugXLVfMMrp8ZSq2tM%3D"
    blob_service_client = BlobServiceClient.from_connection_string(connection_string)
    return blob_service_client

# Function to list folders in the blob storage container
def list_folders():
    blob_service_client = initialize_blob_service_client()
    container_client = blob_service_client.get_container_client(CONTAINER_NAME)

    folders = set()
    for blob in container_client.list_blobs():
        parts = blob.name.split("/")
        if len(parts) > 1:
            # Consider both the root folder and subfolders
            for i in range(1, len(parts)):
                folder_path = "/".join(parts[:i])
                folders.add(folder_path)

    return list(folders)

def load_data():
    with open('C:/Users/mk742457/OneDrive - GSK/one-drive-docs-gsk/2023/DDF-Chatbot/chatbot-webapp/templates/new_flask_app/section.json') as f:
        return json.load(f)

# Initialize sub_sections globally
data = load_data()
sub_sections = [section['sub_section'] for section in data['sections']]

@app.route('/', methods=['GET', 'POST'])
def index():
    try:
        with open('C:/Users/mk742457/OneDrive - GSK/one-drive-docs-gsk/2023/DDF-Chatbot/chatbot-webapp/templates/new_flask_app/section.json') as f:
            data = json.load(f)
        sub_sections = [section['sub_section'] for section in data['sections']]
        return render_template('home.html', folders=list_folders(), summary1=None, sub_sections=sub_sections, selected_sub_section=None, details=None)

    except Exception as e:
        return jsonify({'error': str(e)})

@app.route('/upload', methods=['POST'])
def upload():
    try:
        # Get the selected folder and uploaded file
        selected_folder = request.form.get('selectedFolder')
        file = request.files.get('pdf_file')

        if not selected_folder or not file:
            return jsonify({'error': 'Selected folder or file not provided'})

        # Initialize BlobServiceClient
        blob_service_client = initialize_blob_service_client()
        container_client = blob_service_client.get_container_client(CONTAINER_NAME)

        # Upload the file to Azure Blob Storage
        blob_name = f"{selected_folder}/{file.filename}"
        blob_client = container_client.get_blob_client(blob_name)
        with file.stream as data:
            blob_client.upload_blob(data)

        return jsonify({'message': 'File uploaded successfully'})

    except Exception as e:
        return jsonify({'error': str(e)})

@app.route('/get_details', methods=['POST'])
def get_details():
    selected_sub_section = request.form['sub_section']

    details = None
    for section in data['sections']:
        if section['sub_section'] == selected_sub_section:
            details = section
            break
    return query_entity(details['PartitionKey'], details['RowKey'])

def query_entity(partition_key, row_key):
    my_filter = "PartitionKey eq '"+partition_key+"' and RowKey eq '"+row_key+"'"
    table_client = table_service_client.get_table_client(table_name=table_name)
    entities = table_client.query_entities(my_filter)
    for entity in entities:
        for key in entity.keys():
            if key == "Summary":
                return jsonify({'Summary': entity[key]})

@app.route('/generate_all_summaries', methods=['POST'])
def generate_all_summaries():
    try:
        print("in all summary")
        # Add the code to query entities and create a JSON variable
        query_filter = ""

        # Query entities from the table
        entities = table_client.query_entities(query_filter)

        # Create a list to store the extracted data
        data_list = []

        # Extract required properties and store them in a list of dictionaries
        for entity in entities:
            data_dict = {
                "Section": entity.get("Section", ""),
                "Summary": entity.get("Summary", "")
            }
            data_list.append(data_dict)

        # Return the JSON data directly as a response
        return jsonify(data_list)

    except Exception as e:
        return jsonify({'error': str(e)})

def generatedoc_query_entity(partition_key, row_key):
    txt = ''
    my_filter = "PartitionKey eq '"+partition_key+"' and RowKey eq '"+row_key+"'"
    table_client = table_service_client.get_table_client(table_name=table_name)
    entities = table_client.query_entities(my_filter)
    for entity in entities:
        for key in entity.keys():
            if key == 'Summary':
                txt += str(entity[key])
    return txt

@app.route('/generate_document')
def generate_document():
    try:
        # Load data from the JSON file
        f = open('C:/Users/mk742457/OneDrive - GSK/one-drive-docs-gsk/2023/DDF-Chatbot/chatbot-webapp/templates/new_flask_app/all_file_changed_key_complete.json')
        data_req = json.load(f)
        f.close()

        # Create a Word document
        doc = docx.Document()
        doc.add_picture('gsklogo.jpg', width=Inches(1.25))
        doc.add_heading('Value Evidence Dossier', 0)

        for i in range(0, len(data_req['sections'])):
            main_section_temp = data_req['sections'][i]['PartitionKey']
            section_temp = data_req['sections'][i]['RowKey']
            subsection = data_req['sections'][i]['sub_section_name']
            if(len(generatedoc_query_entity(main_section_temp, section_temp))):
                sub_heading = doc.add_heading(subsection, 1)
                para = doc.add_paragraph().add_run(generatedoc_query_entity(main_section_temp, section_temp))

        # Save the Word document
        doc_filename = "Value_Evidence_Dossier.docx"
        doc.save(doc_filename)

        return send_file(doc_filename, as_attachment=True)

    except Exception as e:
        # Return an error message as JSON
        return jsonify({'success': False, 'error': str(e)})

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
